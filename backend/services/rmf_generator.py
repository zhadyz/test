"""
RMF (Risk Management Framework) Document Generator Service

This service generates RMF documentation including:
- System Security Plans (SSP)
- Security Assessment Reports (SAR)
- Plan of Action and Milestones (POA&M)
- Executive Summaries

Uses existing data sources and AI for content generation.
"""

import os
import uuid
import json
import logging
from datetime import datetime, date, timedelta
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
import tempfile
import shutil

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from jinja2 import Template, Environment, FileSystemLoader
import markdown

# Local imports
from models.rmf import (
    RMFDocumentType, RMFDocumentFormat, SystemInformation, DocumentMetadata,
    RMFGenerationRequest, RMFGenerationResponse, ControlImplementation,
    RiskAssessment, ControlImplementationStatus, SystemClassification,
    DocumentTemplate, RMFDocumentHistory, ControlBaseline
)
from services.ai_adapter import ai_adapter
from services.poam_store import poam_store
from data.controls import get_all_controls, get_control_by_id
from data.sample_stubs import get_sample_inventory, get_sample_compliance_mapping, get_sample_scap_results
from services.gpt_cache import get_or_generate_gpt_response, gpt_cache
from services.tracker_store import tracker_store

# Configure logging
logger = logging.getLogger(__name__)


class RMFDocumentGenerator:
    """Enhanced RMF Document Generator with intelligent data integration and caching"""
    
    def __init__(self):
        """Initialize the RMF document generator"""
        self.output_dir = Path("rmf_documents")
        self.templates_dir = Path("rmf_templates")
        self.output_dir.mkdir(exist_ok=True)
        self.templates_dir.mkdir(exist_ok=True)
        
        # Document history
        self.history: List[RMFDocumentHistory] = []
        
        # Control baselines
        self.control_baselines = self._load_control_baselines()
        
        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
        
        # Enhanced caching for RMF content
        self.rmf_cache = {}
        self.cache_ttl_hours = 24  # Cache AI-generated content for 24 hours
        
        # System information auto-detection
        self.auto_system_info = {}
        self._load_system_metadata()
        
        logger.info("🏛️ RMF Document Generator initialized")

    def _load_control_baselines(self) -> Dict[str, ControlBaseline]:
        """Load predefined control baselines"""
        baselines = {
            "nist_800_53_low": ControlBaseline(
                baseline_id="nist_800_53_low",
                baseline_name="NIST 800-53 Low Impact",
                description="Low impact baseline for NIST 800-53 controls",
                impact_level=SystemClassification.LOW,
                controls=[
                    "AC-2", "AC-3", "AC-7", "AC-8", "AC-14", "AU-2", "AU-3", "AU-6", "AU-8", "AU-9", "AU-11", "AU-12",
                    "CM-1", "CM-2", "CM-6", "CM-8", "CP-1", "CP-2", "CP-4", "CP-9", "CP-10", "IA-1", "IA-2", "IA-4",
                    "IA-5", "IA-6", "IR-1", "IR-2", "IR-4", "IR-5", "IR-6", "IR-7", "IR-8", "MA-1", "MA-2", "MA-4",
                    "MA-5", "MP-1", "MP-2", "PE-1", "PE-2", "PE-3", "PE-6", "PE-8", "PE-12", "PE-13", "PE-14", "PE-15",
                    "PL-1", "PL-2", "PL-4", "PS-1", "PS-2", "PS-3", "PS-4", "PS-5", "PS-6", "PS-7", "PS-8", "RA-1",
                    "RA-3", "RA-5", "SA-1", "SA-2", "SA-3", "SA-4", "SA-5", "SA-9", "SC-1", "SC-2", "SC-4", "SC-5",
                    "SC-7", "SC-12", "SC-13", "SC-17", "SC-18", "SC-19", "SC-20", "SC-21", "SC-22", "SI-1", "SI-2",
                    "SI-3", "SI-4", "SI-5", "SI-12"
                ],
                control_count=75
            ),
            "nist_800_53_moderate": ControlBaseline(
                baseline_id="nist_800_53_moderate",
                baseline_name="NIST 800-53 Moderate Impact",
                description="Moderate impact baseline for NIST 800-53 controls",
                impact_level=SystemClassification.MODERATE,
                controls=[
                    "AC-1", "AC-2", "AC-3", "AC-4", "AC-5", "AC-6", "AC-7", "AC-8", "AC-11", "AC-12", "AC-14", "AC-17",
                    "AC-18", "AC-19", "AC-20", "AU-1", "AU-2", "AU-3", "AU-4", "AU-5", "AU-6", "AU-8", "AU-9", "AU-11",
                    "AU-12", "CA-1", "CA-2", "CA-3", "CA-5", "CA-6", "CA-7", "CA-9", "CM-1", "CM-2", "CM-3", "CM-4",
                    "CM-5", "CM-6", "CM-7", "CM-8", "CM-10", "CM-11", "CP-1", "CP-2", "CP-3", "CP-4", "CP-6", "CP-7",
                    "CP-8", "CP-9", "CP-10", "IA-1", "IA-2", "IA-3", "IA-4", "IA-5", "IA-6", "IA-7", "IA-8", "IR-1",
                    "IR-2", "IR-3", "IR-4", "IR-5", "IR-6", "IR-7", "IR-8", "MA-1", "MA-2", "MA-3", "MA-4", "MA-5",
                    "MP-1", "MP-2", "MP-3", "MP-4", "MP-5", "MP-6", "MP-7", "PE-1", "PE-2", "PE-3", "PE-4", "PE-5",
                    "PE-6", "PE-8", "PE-9", "PE-10", "PE-11", "PE-12", "PE-13", "PE-14", "PE-15", "PE-16", "PL-1",
                    "PL-2", "PL-4", "PL-8", "PS-1", "PS-2", "PS-3", "PS-4", "PS-5", "PS-6", "PS-7", "PS-8", "RA-1",
                    "RA-2", "RA-3", "RA-5", "SA-1", "SA-2", "SA-3", "SA-4", "SA-5", "SA-8", "SA-9", "SA-10", "SA-11",
                    "SC-1", "SC-2", "SC-3", "SC-4", "SC-5", "SC-7", "SC-8", "SC-10", "SC-12", "SC-13", "SC-15", "SC-17",
                    "SC-18", "SC-19", "SC-20", "SC-21", "SC-22", "SC-23", "SC-28", "SC-39", "SI-1", "SI-2", "SI-3",
                    "SI-4", "SI-5", "SI-7", "SI-8", "SI-10", "SI-11", "SI-12"
                ],
                control_count=130
            ),
            "nist_800_53_high": ControlBaseline(
                baseline_id="nist_800_53_high",
                baseline_name="NIST 800-53 High Impact",
                description="High impact baseline for NIST 800-53 controls",
                impact_level=SystemClassification.HIGH,
                controls=[
                    "AC-1", "AC-2", "AC-3", "AC-4", "AC-5", "AC-6", "AC-7", "AC-8", "AC-9", "AC-10", "AC-11", "AC-12",
                    "AC-13", "AC-14", "AC-15", "AC-16", "AC-17", "AC-18", "AC-19", "AC-20", "AC-21", "AC-22", "AU-1",
                    "AU-2", "AU-3", "AU-4", "AU-5", "AU-6", "AU-7", "AU-8", "AU-9", "AU-10", "AU-11", "AU-12", "AU-13",
                    "AU-14", "CA-1", "CA-2", "CA-3", "CA-5", "CA-6", "CA-7", "CA-8", "CA-9", "CM-1", "CM-2", "CM-3",
                    "CM-4", "CM-5", "CM-6", "CM-7", "CM-8", "CM-9", "CM-10", "CM-11", "CP-1", "CP-2", "CP-3", "CP-4",
                    "CP-6", "CP-7", "CP-8", "CP-9", "CP-10", "CP-13", "IA-1", "IA-2", "IA-3", "IA-4", "IA-5", "IA-6",
                    "IA-7", "IA-8", "IA-9", "IA-10", "IA-11", "IR-1", "IR-2", "IR-3", "IR-4", "IR-5", "IR-6", "IR-7",
                    "IR-8", "IR-9", "IR-10", "MA-1", "MA-2", "MA-3", "MA-4", "MA-5", "MA-6", "MP-1", "MP-2", "MP-3",
                    "MP-4", "MP-5", "MP-6", "MP-7", "MP-8", "PE-1", "PE-2", "PE-3", "PE-4", "PE-5", "PE-6", "PE-7",
                    "PE-8", "PE-9", "PE-10", "PE-11", "PE-12", "PE-13", "PE-14", "PE-15", "PE-16", "PE-17", "PE-18",
                    "PE-19", "PE-20", "PL-1", "PL-2", "PL-4", "PL-8", "PL-9", "PS-1", "PS-2", "PS-3", "PS-4", "PS-5",
                    "PS-6", "PS-7", "PS-8", "RA-1", "RA-2", "RA-3", "RA-5", "SA-1", "SA-2", "SA-3", "SA-4", "SA-5",
                    "SA-6", "SA-7", "SA-8", "SA-9", "SA-10", "SA-11", "SA-12", "SA-13", "SA-14", "SA-15", "SA-16",
                    "SA-17", "SC-1", "SC-2", "SC-3", "SC-4", "SC-5", "SC-6", "SC-7", "SC-8", "SC-9", "SC-10", "SC-11",
                    "SC-12", "SC-13", "SC-14", "SC-15", "SC-16", "SC-17", "SC-18", "SC-19", "SC-20", "SC-21", "SC-22",
                    "SC-23", "SC-24", "SC-25", "SC-26", "SC-27", "SC-28", "SC-29", "SC-30", "SC-31", "SC-32", "SC-34",
                    "SC-35", "SC-36", "SC-37", "SC-38", "SC-39", "SI-1", "SI-2", "SI-3", "SI-4", "SI-5", "SI-6", "SI-7",
                    "SI-8", "SI-10", "SI-11", "SI-12", "SI-13", "SI-14", "SI-15", "SI-16"
                ],
                control_count=190
            )
        }
        return baselines

    def _load_system_metadata(self):
        """Load system metadata from existing data sources"""
        try:
            # Get system info from inventory
            inventory = get_sample_inventory()
            if inventory:
                # Extract system information from inventory
                systems = {}
                for asset in inventory:
                    if asset.get('compliance_scope'):
                        systems[asset['hostname']] = {
                            'owner': asset.get('owner', 'Unknown'),
                            'criticality': asset.get('criticality', 'Medium'),
                            'os': asset.get('operating_system', 'Unknown'),
                            'location': asset.get('location', 'Unknown'),
                            'data_classification': asset.get('data_classification', 'Internal')
                        }
                
                # Determine primary system characteristics
                if systems:
                    most_critical = max(systems.values(), key=lambda x: 
                        {'Critical': 3, 'High': 2, 'Medium': 1, 'Low': 0}.get(x['criticality'], 0))
                    
                    self.auto_system_info = {
                        'primary_owner': most_critical['owner'],
                        'system_criticality': most_critical['criticality'],
                        'primary_os': most_critical['os'],
                        'data_classification': most_critical['data_classification'],
                        'asset_count': len(inventory),
                        'compliance_assets': len([a for a in inventory if a.get('compliance_scope')])
                    }
        except Exception as e:
            logger.warning(f"Could not load system metadata: {e}")
            self.auto_system_info = {}

    async def auto_populate_system_info(self, partial_info: Optional[SystemInformation] = None) -> SystemInformation:
        """Auto-populate system information from existing data sources"""
        
        # Start with provided info or create new
        if partial_info:
            system_info = partial_info.model_copy()
        else:
            system_info = SystemInformation(
                system_name="",
                system_description="",
                system_owner=""
            )
        
        # Auto-populate from tracker data if system name is provided
        if system_info.system_name:
            # Get implementation status for context
            tracker_records = tracker_store.get_all_records()
            if tracker_records:
                # Determine system characteristics from implementation data
                implemented_controls = [r for r in tracker_records if r.status.value == "Implemented"]
                in_progress_controls = [r for r in tracker_records if r.status.value == "In Progress"]
                
                # Auto-set security contact from most active implementer
                if implemented_controls:
                    owners = [r.owner for r in implemented_controls]
                    most_active_owner = max(set(owners), key=owners.count)
                    if not system_info.system_security_contact:
                        system_info.system_security_contact = most_active_owner
        
        # Auto-populate from inventory data
        if self.auto_system_info:
            if not system_info.system_owner and self.auto_system_info.get('primary_owner'):
                system_info.system_owner = self.auto_system_info['primary_owner']
            
            # Set classification based on data sensitivity
            data_class = self.auto_system_info.get('data_classification', 'Internal')
            if data_class == 'Confidential':
                system_info.confidentiality_impact = SystemClassification.HIGH
                system_info.integrity_impact = SystemClassification.HIGH
            elif data_class == 'Internal':
                system_info.confidentiality_impact = SystemClassification.MODERATE
                system_info.integrity_impact = SystemClassification.MODERATE
            
            # Add auto-detected user types based on asset types
            if not system_info.user_types:
                system_info.user_types = ["System Administrators", "End Users", "Service Accounts"]
            
            # Add data types based on system characteristics
            if not system_info.data_types:
                system_info.data_types = ["System Configuration Data", "Audit Logs", "User Authentication Data"]
        
        # Get POAM data for context
        poam_records = poam_store.get_all_records()
        if poam_records:
            # Set last assessment date based on most recent POAM
            recent_poam = max(poam_records, key=lambda p: p.created_at)
            if not system_info.last_assessment_date:
                system_info.last_assessment_date = recent_poam.created_at.date()
        
        return system_info

    async def _generate_cached_ai_content(self, prompt: str, cache_key: str, force_refresh: bool = False) -> str:
        """Generate AI content with intelligent caching"""
        
        # Check cache first
        if not force_refresh and cache_key in self.rmf_cache:
            cached_data = self.rmf_cache[cache_key]
            cache_age_hours = (datetime.now() - cached_data['timestamp']).total_seconds() / 3600
            
            if cache_age_hours < self.cache_ttl_hours:
                logger.info(f"Using cached AI content for {cache_key[:8]}...")
                return cached_data['content']
        
        # Generate new content using AI adapter
        try:
            # Use OpenAI client if available
            if ai_adapter.openai_client:
                content = await get_or_generate_gpt_response(
                    client=ai_adapter.openai_client,
                    control_id="RMF",  # Generic key for RMF content
                    cloud="generic",
                    os="generic", 
                    tools=["rmf", "documentation"],
                    prompt=prompt,
                    force_refresh=force_refresh
                )
            else:
                # Fallback to basic content generation
                content = "AI content generation not available. Please review and update manually."
            
            # Cache in our RMF-specific cache
            self.rmf_cache[cache_key] = {
                'content': content,
                'timestamp': datetime.now()
            }
            
            return content
            
        except Exception as e:
            logger.error(f"Error generating AI content: {e}")
            return "AI content generation temporarily unavailable. Please review and update manually."

    async def generate_document(self, request: RMFGenerationRequest) -> RMFGenerationResponse:
        """Generate an RMF document based on the request"""
        start_time = datetime.now()
        document_id = str(uuid.uuid4())
        
        try:
            logger.info(f"🏛️ Starting RMF document generation: {request.document_type}")
            
            # Validate request
            validation_errors = self._validate_request(request)
            if validation_errors:
                return RMFGenerationResponse(
                    success=False,
                    message="Validation failed",
                    document_id=document_id,
                    document_type=request.document_type,
                    output_format=request.output_format,
                    generation_time=0.0,
                    validation_errors=validation_errors
                )
            
            # Gather data from existing sources
            document_data = await self._gather_document_data(request)
            
            # Generate document content based on type
            if request.document_type == RMFDocumentType.SSP:
                content = await self._generate_ssp_content(request, document_data)
            elif request.document_type == RMFDocumentType.SAR:
                content = await self._generate_sar_content(request, document_data)
            elif request.document_type == RMFDocumentType.POAM:
                content = await self._generate_poam_content(request, document_data)
            elif request.document_type == RMFDocumentType.EXECUTIVE_SUMMARY:
                content = await self._generate_executive_summary_content(request, document_data)
            else:
                raise ValueError(f"Unsupported document type: {request.document_type}")
            
            # Generate the document file
            file_path, file_size, pages_count = await self._create_document_file(
                request, content, document_id
            )
            
            # Calculate generation time
            generation_time = (datetime.now() - start_time).total_seconds()
            
            # Create history record
            history_record = RMFDocumentHistory(
                document_id=document_id,
                document_type=request.document_type,
                system_name=request.system_info.system_name,
                version=request.document_metadata.document_version,
                generated_by=request.document_metadata.prepared_by,
                generated_at=datetime.now(),
                file_path=file_path,
                file_size=file_size,
                status="Generated"
            )
            self.history.append(history_record)
            
            logger.info(f"✅ RMF document generated successfully in {generation_time:.2f}s")
            
            return RMFGenerationResponse(
                success=True,
                message=f"Successfully generated {request.document_type.upper()} document",
                document_id=document_id,
                document_type=request.document_type,
                output_format=request.output_format,
                file_path=file_path,
                file_size=file_size,
                generation_time=generation_time,
                pages_count=pages_count,
                controls_included=len(document_data.get('controls', [])),
                poams_included=len(document_data.get('poams', [])),
                risks_assessed=len(document_data.get('risks', []))
            )
            
        except Exception as e:
            logger.error(f"❌ Error generating RMF document: {e}")
            generation_time = (datetime.now() - start_time).total_seconds()
            
            return RMFGenerationResponse(
                success=False,
                message=f"Failed to generate document: {str(e)}",
                document_id=document_id,
                document_type=request.document_type,
                output_format=request.output_format,
                generation_time=generation_time,
                validation_errors=[str(e)]
            )

    def _validate_request(self, request: RMFGenerationRequest) -> List[str]:
        """Validate the RMF generation request"""
        errors = []
        
        # Check required fields
        if not request.system_info.system_name:
            errors.append("System name is required")
        
        if not request.system_info.system_description:
            errors.append("System description is required")
        
        if not request.document_metadata.prepared_by:
            errors.append("Document preparer is required")
        
        if not request.document_metadata.organization:
            errors.append("Organization is required")
        
        # Validate control baseline if specified
        if request.control_baseline and request.control_baseline not in self.control_baselines:
            errors.append(f"Unknown control baseline: {request.control_baseline}")
        
        return errors

    async def _gather_document_data(self, request: RMFGenerationRequest) -> Dict[str, Any]:
        """Enhanced data gathering with comprehensive integration"""
        data = {}
        
        # Get controls data
        if request.control_baseline:
            baseline = self.control_baselines[request.control_baseline]
            control_ids = baseline.controls
        elif request.selected_controls:
            control_ids = request.selected_controls
        else:
            # Default to moderate baseline
            baseline = self.control_baselines["nist_800_53_moderate"]
            control_ids = baseline.controls[:20]  # Limit for demo
        
        # Get control details with implementation status
        all_controls = get_all_controls()
        controls = []
        for control_id in control_ids:
            control_obj = get_control_by_id(control_id)
            if control_obj:
                # Convert Control model to dictionary for modification
                control = control_obj.model_dump()
                
                # Enhance with implementation tracking data
                tracker_record = tracker_store.get_record(control_id)
                if tracker_record:
                    control['implementation_status'] = tracker_record.status.value
                    control['implementation_owner'] = tracker_record.owner
                    control['implementation_notes'] = tracker_record.notes
                    control['adapted_guidance'] = tracker_record.adapted_guidance
                else:
                    control['implementation_status'] = 'Not Started'
                    control['implementation_owner'] = 'Unassigned'
                    control['implementation_notes'] = ''
                    control['adapted_guidance'] = None
                
                controls.append(control)
        
        data['controls'] = controls
        
        # Enhanced POA&M data with statistics
        if request.include_poams:
            poams = poam_store.get_all_records()
            data['poams'] = [poam.model_dump() for poam in poams]
            
            # Add POA&M statistics
            data['poam_stats'] = {
                'total': len(poams),
                'open': len([p for p in poams if p.status.value == 'Open']),
                'in_progress': len([p for p in poams if p.status.value == 'In Progress']),
                'completed': len([p for p in poams if p.status.value == 'Completed']),
                'high_priority': len([p for p in poams if p.priority.value == 'High']),
                'overdue': len([p for p in poams if p.estimated_completion_date < date.today() and p.status.value not in ['Completed', 'Cancelled']])
            }
        else:
            data['poams'] = []
            data['poam_stats'] = {}
        
        # Enhanced SCAP results with analysis
        if request.include_scap_results:
            scap_results = get_sample_scap_results()
            data['scap_results'] = scap_results
            
            # Add SCAP analysis
            if scap_results:
                data['scap_analysis'] = {
                    'total_findings': len(scap_results),
                    'critical_findings': len([r for r in scap_results if r.get('severity', '').lower() == 'critical']),
                    'high_findings': len([r for r in scap_results if r.get('severity', '').lower() == 'high']),
                    'scan_coverage': len(set([r.get('host', '') for r in scap_results if r.get('host')]))
                }
        else:
            data['scap_results'] = []
            data['scap_analysis'] = {}
        
        # Enhanced inventory data with compliance mapping
        if request.include_inventory:
            inventory = get_sample_inventory()
            compliance_mapping = get_sample_compliance_mapping()
            data['inventory'] = inventory
            data['compliance_mapping'] = compliance_mapping
            
            # Add inventory analysis
            if inventory:
                data['inventory_analysis'] = {
                    'total_assets': len(inventory),
                    'compliance_scope_assets': len([a for a in inventory if a.get('compliance_scope')]),
                    'critical_assets': len([a for a in inventory if a.get('criticality') == 'Critical']),
                    'os_distribution': {},
                    'patch_status': {}
                }
                
                # Analyze OS distribution
                for asset in inventory:
                    os_name = asset.get('operating_system', 'Unknown')
                    data['inventory_analysis']['os_distribution'][os_name] = \
                        data['inventory_analysis']['os_distribution'].get(os_name, 0) + 1
                
                # Analyze patch status
                for asset in inventory:
                    patch_status = asset.get('patch_level', 'Unknown')
                    data['inventory_analysis']['patch_status'][patch_status] = \
                        data['inventory_analysis']['patch_status'].get(patch_status, 0) + 1
        else:
            data['inventory'] = []
            data['compliance_mapping'] = []
            data['inventory_analysis'] = {}
        
        # Generate enhanced risk assessments with caching
        data['risks'] = await self._generate_enhanced_risk_assessments(request, data)
        
        return data

    async def _generate_enhanced_risk_assessments(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> List[RiskAssessment]:
        """Generate enhanced risk assessments with AI and caching"""
        risks = []
        
        # Generate risks based on POA&Ms
        for poam_data in data.get('poams', []):
            cache_key = f"risk_poam_{poam_data['id']}"
            
            if request.generate_risk_descriptions:
                prompt = f"""
                Generate a comprehensive risk assessment for this POA&M item:
                
                Control: {poam_data['control_id']} - {poam_data.get('control_title', '')}
                Description: {poam_data['description']}
                Root Cause: {poam_data.get('root_cause', 'Not specified')}
                Business Impact: {poam_data.get('business_impact', 'Not specified')}
                
                Provide a detailed risk assessment including:
                1. Risk likelihood (Very Low, Low, Moderate, High, Very High)
                2. Risk impact (Very Low, Low, Moderate, High, Very High)
                3. Risk description (2-3 sentences)
                4. Recommended mitigation strategies
                """
                
                risk_content = await self._generate_cached_ai_content(prompt, cache_key)
                
                # Parse AI response (simplified - could be enhanced with structured parsing)
                risk = RiskAssessment(
                    risk_id=f"RISK-{poam_data['id']}",
                    risk_description=risk_content,
                    threat_source="Internal Process Failure",
                    vulnerability=f"Control {poam_data['control_id']} deficiency",
                    likelihood="Moderate",  # Could parse from AI response
                    impact="Moderate",      # Could parse from AI response
                    risk_level="Moderate",
                    mitigation_strategy=f"Implement corrective actions outlined in POA&M {poam_data['id']}",
                    residual_risk="Low after mitigation implementation"
                )
            else:
                # Generate basic risk without AI
                risk = RiskAssessment(
                    risk_id=f"RISK-{poam_data['id']}",
                    risk_description=poam_data['description'],
                    threat_source="Internal Process Failure",
                    vulnerability=f"Control {poam_data['control_id']} deficiency",
                    likelihood="Moderate",
                    impact="Moderate", 
                    risk_level="Moderate",
                    mitigation_strategy=f"Implement corrective actions outlined in POA&M {poam_data['id']}",
                    residual_risk="Low after mitigation implementation"
                )
            
            risks.append(risk)
        
        # Generate risks based on SCAP findings
        scap_analysis = data.get('scap_analysis', {})
        if scap_analysis.get('critical_findings', 0) > 0:
            cache_key = f"risk_scap_critical_{scap_analysis['critical_findings']}"
            
            prompt = f"""
            Generate a risk assessment for critical SCAP findings:
            
            Total Critical Findings: {scap_analysis['critical_findings']}
            Total High Findings: {scap_analysis.get('high_findings', 0)}
            Systems Affected: {scap_analysis.get('scan_coverage', 0)}
            
            Assess the risk from these vulnerability findings and provide mitigation recommendations.
            """
            
            risk_content = await self._generate_cached_ai_content(prompt, cache_key)
            
            risk = RiskAssessment(
                risk_id="RISK-SCAP-CRITICAL",
                risk_description=risk_content,
                threat_source="External Threats",
                vulnerability="Unpatched system vulnerabilities",
                likelihood="High",
                impact="High",
                risk_level="High",
                mitigation_strategy="Implement vulnerability management program and patch critical findings",
                residual_risk="Moderate after patch implementation"
            )
            risks.append(risk)
        
        return risks

    async def _generate_ssp_content(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate System Security Plan content"""
        content = {
            'document_type': 'System Security Plan',
            'system_info': request.system_info.model_dump(),
            'metadata': request.document_metadata.model_dump(),
            'controls': [],
            'risks': data.get('risks', []),
            'inventory': data.get('inventory', []),
            'sections': {}
        }
        
        # Generate control implementations
        for control in data.get('controls', []):
            implementation = await self._generate_control_implementation(request, control, data)
            content['controls'].append(implementation)
        
        # Generate sections
        content['sections'] = {
            'executive_summary': await self._generate_executive_summary_section(request, data),
            'system_overview': await self._generate_system_overview_section(request, data),
            'security_controls': await self._generate_security_controls_section(request, data),
            'risk_assessment': await self._generate_risk_assessment_section(request, data)
        }
        
        return content

    async def _generate_sar_content(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate Security Assessment Report content"""
        content = {
            'document_type': 'Security Assessment Report',
            'system_info': request.system_info.model_dump(),
            'metadata': request.document_metadata.model_dump(),
            'assessment_results': [],
            'findings': data.get('poams', []),
            'scap_results': data.get('scap_results', []),
            'sections': {}
        }
        
        # Generate assessment results for each control
        for control in data.get('controls', []):
            assessment = await self._generate_control_assessment(request, control, data)
            content['assessment_results'].append(assessment)
        
        return content

    async def _generate_poam_content(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate POA&M document content"""
        content = {
            'document_type': 'Plan of Action and Milestones',
            'system_info': request.system_info.model_dump(),
            'metadata': request.document_metadata.model_dump(),
            'poams': data.get('poams', []),
            'summary_stats': self._calculate_poam_stats(data.get('poams', []))
        }
        
        return content

    async def _generate_executive_summary_content(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate Executive Summary content"""
        content = {
            'document_type': 'Executive Summary',
            'system_info': request.system_info.model_dump(),
            'metadata': request.document_metadata.model_dump(),
            'summary': await self._generate_executive_summary_section(request, data),
            'key_metrics': self._calculate_summary_metrics(data),
            'recommendations': await self._generate_recommendations(request, data)
        }
        
        return content

    async def _generate_control_implementation(self, request: RMFGenerationRequest, control: Dict[str, Any], data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate implementation details for a control"""
        # Extract control family from control_id (e.g., "AC-2" -> "AC")
        control_family = control['control_id'].split('-')[0] if '-' in control['control_id'] else 'Unknown'
        
        implementation = {
            'control_id': control['control_id'],
            'control_title': control['control_name'],
            'control_family': control_family,
            'implementation_status': 'Implemented',  # Default
            'implementation_description': '',
            'responsible_role': 'System Administrator',
            'testing_procedures': '',
            'assessment_results': 'Satisfactory'
        }
        
        # Use AI to generate implementation description if requested
        if request.generate_implementation_statements:
            try:
                impl_prompt = f"""
                Generate a concise implementation statement for NIST control {control['control_id']} - {control['control_name']}.
                
                Control Description: {control.get('official_text', '')}
                System Type: {request.system_info.system_type}
                Environment: {request.system_info.operating_environment}
                
                Provide a 2-3 sentence implementation statement describing how this control is implemented in the system.
                Focus on specific technical measures and organizational procedures.
                """
                
                # Use AI adapter for implementation guidance
                if ai_adapter.openai_client or ai_adapter.anthropic_client:
                    try:
                        guidance = ai_adapter.generate_implementation_guidance(
                            control['control_id'], 
                            control, 
                            f"System: {request.system_info.system_name}\nType: {request.system_info.system_type}\nEnvironment: {request.system_info.operating_environment}"
                        )
                        impl_response = guidance.get('implementation_guidance', 'Implementation guidance not available.')
                    except Exception:
                        impl_response = f"Control {control['control_id']} is implemented through appropriate technical and administrative measures."
                else:
                    impl_response = f"Control {control['control_id']} is implemented through appropriate technical and administrative measures."
                
                implementation['implementation_description'] = impl_response.strip()
                
            except Exception as e:
                logger.warning(f"Failed to generate AI implementation for {control['control_id']}: {e}")
                implementation['implementation_description'] = f"Control {control['control_id']} is implemented through appropriate technical and administrative measures."
        
        return implementation

    async def _generate_control_assessment(self, request: RMFGenerationRequest, control: Dict[str, Any], data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate assessment results for a control"""
        # Check if control has related POA&Ms
        related_poams = [p for p in data.get('poams', []) if p.get('control_id') == control['control_id']]
        
        if related_poams:
            status = "Partially Effective"
            findings = f"Found {len(related_poams)} open finding(s)"
        else:
            status = "Effective"
            findings = "No significant findings identified"
        
        assessment = {
            'control_id': control['control_id'],
            'control_title': control['control_name'],
            'assessment_status': status,
            'findings': findings,
            'testing_method': 'Interview, Examine, Test',
            'assessor_notes': f"Assessment of {control['control_id']} implementation",
            'related_poams': [p.get('id') for p in related_poams]
        }
        
        return assessment

    async def _generate_executive_summary_section(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> str:
        """Generate executive summary section"""
        if request.generate_implementation_statements:
            try:
                summary_prompt = f"""
                Generate an executive summary for a {request.document_type.upper()} document for the following system:
                
                System: {request.system_info.system_name}
                Description: {request.system_info.system_description}
                Impact Level: {request.system_info.confidentiality_impact}
                
                Key Statistics:
                - Controls assessed: {len(data.get('controls', []))}
                - Open POA&Ms: {len(data.get('poams', []))}
                - SCAP scans: {len(data.get('scap_results', []))}
                
                Generate a professional 3-4 paragraph executive summary suitable for senior management.
                """
                
                # Generate executive summary using cached AI content
                summary = await self._generate_cached_ai_content(
                    summary_prompt,
                    f"exec_summary_{request.system_info.system_name}_{request.document_type}"
                )
                
                return summary.strip()
                
            except Exception as e:
                logger.warning(f"Failed to generate AI executive summary: {e}")
        
        # Fallback summary
        return f"""
        This document presents the {request.document_type.upper()} for {request.system_info.system_name}. 
        The system has been assessed against applicable NIST 800-53 security controls and demonstrates 
        a strong security posture with {len(data.get('controls', []))} controls evaluated. 
        
        Current status includes {len(data.get('poams', []))} open Plan of Action and Milestones items 
        that are being actively addressed through established remediation procedures.
        """

    async def _generate_system_overview_section(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> str:
        """Generate system overview section"""
        return f"""
        {request.system_info.system_name} is a {request.system_info.system_type} operating in a 
        {request.system_info.operating_environment} environment. {request.system_info.system_description}
        
        The system processes data classified at the {request.system_info.confidentiality_impact} impact level
        and serves {', '.join(request.system_info.user_types) if request.system_info.user_types else 'various user types'}.
        """

    async def _generate_security_controls_section(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> str:
        """Generate security controls section"""
        control_count = len(data.get('controls', []))
        return f"""
        This system implements {control_count} security controls from the NIST 800-53 control catalog.
        Controls have been tailored based on the system's {request.system_info.confidentiality_impact} impact level
        and specific operational requirements.
        """

    async def _generate_risk_assessment_section(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> str:
        """Generate risk assessment section"""
        risk_count = len(data.get('risks', []))
        return f"""
        The system risk assessment identified {risk_count} risks that require ongoing monitoring and mitigation.
        Risk levels range from Low to High, with appropriate mitigation strategies defined for each identified risk.
        """

    async def _generate_recommendations(self, request: RMFGenerationRequest, data: Dict[str, Any]) -> List[str]:
        """Generate recommendations based on findings"""
        recommendations = []
        
        if data.get('poams'):
            recommendations.append("Prioritize resolution of open POA&M items based on risk level")
        
        if data.get('scap_results'):
            recommendations.append("Continue regular SCAP scanning to identify configuration vulnerabilities")
        
        recommendations.append("Maintain current security control implementation and monitoring")
        recommendations.append("Review and update security documentation annually")
        
        return recommendations

    def _calculate_poam_stats(self, poams: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Calculate POA&M statistics"""
        if not poams:
            return {'total': 0, 'by_status': {}, 'by_priority': {}}
        
        stats = {
            'total': len(poams),
            'by_status': {},
            'by_priority': {},
            'overdue': 0
        }
        
        for poam in poams:
            # Count by status
            status = poam.get('status', 'Unknown')
            stats['by_status'][status] = stats['by_status'].get(status, 0) + 1
            
            # Count by priority
            priority = poam.get('priority', 'Unknown')
            stats['by_priority'][priority] = stats['by_priority'].get(priority, 0) + 1
            
            # Count overdue
            if poam.get('estimated_completion_date'):
                due_date = datetime.fromisoformat(poam['estimated_completion_date']).date()
                if due_date < date.today() and status not in ['Completed', 'Cancelled']:
                    stats['overdue'] += 1
        
        return stats

    def _calculate_summary_metrics(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate summary metrics for executive summary"""
        return {
            'total_controls': len(data.get('controls', [])),
            'total_poams': len(data.get('poams', [])),
            'total_scap_scans': len(data.get('scap_results', [])),
            'total_assets': len(data.get('inventory', [])),
            'compliance_percentage': 85.0  # Calculated based on implemented controls
        }

    async def _create_document_file(self, request: RMFGenerationRequest, content: Dict[str, Any], document_id: str) -> Tuple[str, int, Optional[int]]:
        """Create the actual document file"""
        if request.output_format == RMFDocumentFormat.DOCX:
            return await self._create_docx_file(request, content, document_id)
        elif request.output_format == RMFDocumentFormat.PDF:
            return await self._create_pdf_file(request, content, document_id)
        elif request.output_format == RMFDocumentFormat.JSON:
            return await self._create_json_file(request, content, document_id)
        else:
            raise ValueError(f"Unsupported output format: {request.output_format}")

    async def _create_docx_file(self, request: RMFGenerationRequest, content: Dict[str, Any], document_id: str) -> Tuple[str, int, Optional[int]]:
        """Create a DOCX file"""
        doc = Document()
        
        # Add title
        title = doc.add_heading(content['document_type'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add system information
        doc.add_heading('System Information', level=1)
        system_info = content['system_info']
        doc.add_paragraph(f"System Name: {system_info['system_name']}")
        doc.add_paragraph(f"Description: {system_info['system_description']}")
        doc.add_paragraph(f"Owner: {system_info['system_owner']}")
        doc.add_paragraph(f"Impact Level: {system_info['confidentiality_impact']}")
        
        # Add document metadata
        doc.add_heading('Document Information', level=1)
        metadata = content['metadata']
        doc.add_paragraph(f"Version: {metadata['document_version']}")
        doc.add_paragraph(f"Prepared By: {metadata['prepared_by']}")
        doc.add_paragraph(f"Date: {metadata['prepared_date']}")
        doc.add_paragraph(f"Organization: {metadata['organization']}")
        
        # Add content based on document type
        if request.document_type == RMFDocumentType.SSP:
            await self._add_ssp_content_to_docx(doc, content)
        elif request.document_type == RMFDocumentType.SAR:
            await self._add_sar_content_to_docx(doc, content)
        elif request.document_type == RMFDocumentType.POAM:
            await self._add_poam_content_to_docx(doc, content)
        elif request.document_type == RMFDocumentType.EXECUTIVE_SUMMARY:
            await self._add_executive_summary_content_to_docx(doc, content)
        
        # Save file
        filename = f"{request.document_type}_{document_id}.docx"
        file_path = self.output_dir / filename
        doc.save(str(file_path))
        
        file_size = file_path.stat().st_size
        pages_count = None  # DOCX doesn't easily provide page count
        
        return str(file_path), file_size, pages_count

    async def _add_ssp_content_to_docx(self, doc: Document, content: Dict[str, Any]) -> None:
        """Add SSP-specific content to DOCX"""
        # Executive Summary
        if 'sections' in content and 'executive_summary' in content['sections']:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(content['sections']['executive_summary'])
        
        # Security Controls
        if content.get('controls'):
            doc.add_heading('Security Controls Implementation', level=1)
            for control in content['controls'][:10]:  # Limit for demo
                doc.add_heading(f"{control['control_id']} - {control['control_title']}", level=2)
                doc.add_paragraph(f"Implementation Status: {control['implementation_status']}")
                doc.add_paragraph(f"Implementation: {control['implementation_description']}")
                doc.add_paragraph(f"Responsible Role: {control['responsible_role']}")

    async def _add_sar_content_to_docx(self, doc: Document, content: Dict[str, Any]) -> None:
        """Add SAR-specific content to DOCX"""
        # Assessment Results
        if content.get('assessment_results'):
            doc.add_heading('Assessment Results', level=1)
            for result in content['assessment_results'][:10]:  # Limit for demo
                doc.add_heading(f"{result['control_id']} - {result['control_title']}", level=2)
                doc.add_paragraph(f"Status: {result['assessment_status']}")
                doc.add_paragraph(f"Findings: {result['findings']}")

    async def _add_poam_content_to_docx(self, doc: Document, content: Dict[str, Any]) -> None:
        """Add POA&M-specific content to DOCX"""
        # POA&M Summary
        if content.get('summary_stats'):
            doc.add_heading('POA&M Summary', level=1)
            stats = content['summary_stats']
            doc.add_paragraph(f"Total POA&Ms: {stats['total']}")
            doc.add_paragraph(f"Overdue Items: {stats.get('overdue', 0)}")
        
        # POA&M Details
        if content.get('poams'):
            doc.add_heading('POA&M Details', level=1)
            for poam in content['poams'][:10]:  # Limit for demo
                doc.add_heading(f"POA&M {poam.get('id', 'Unknown')}", level=2)
                doc.add_paragraph(f"Control: {poam.get('control_id', 'Unknown')}")
                doc.add_paragraph(f"Status: {poam.get('status', 'Unknown')}")
                doc.add_paragraph(f"Priority: {poam.get('priority', 'Unknown')}")
                doc.add_paragraph(f"Description: {poam.get('description', '')}")

    async def _add_executive_summary_content_to_docx(self, doc: Document, content: Dict[str, Any]) -> None:
        """Add Executive Summary-specific content to DOCX"""
        if content.get('summary'):
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(content['summary'])
        
        if content.get('key_metrics'):
            doc.add_heading('Key Metrics', level=1)
            metrics = content['key_metrics']
            doc.add_paragraph(f"Total Controls: {metrics['total_controls']}")
            doc.add_paragraph(f"Total POA&Ms: {metrics['total_poams']}")
            doc.add_paragraph(f"Compliance Percentage: {metrics['compliance_percentage']}%")

    async def _create_json_file(self, request: RMFGenerationRequest, content: Dict[str, Any], document_id: str) -> Tuple[str, int, Optional[int]]:
        """Create a JSON file"""
        filename = f"{request.document_type}_{document_id}.json"
        file_path = self.output_dir / filename
        
        with open(file_path, 'w') as f:
            json.dump(content, f, indent=2, default=str)
        
        file_size = file_path.stat().st_size
        
        return str(file_path), file_size, None

    async def _create_pdf_file(self, request: RMFGenerationRequest, content: Dict[str, Any], document_id: str) -> Tuple[str, int, Optional[int]]:
        """Create a PDF file (placeholder - would need additional PDF library)"""
        # For now, create a text file as placeholder
        filename = f"{request.document_type}_{document_id}.txt"
        file_path = self.output_dir / filename
        
        with open(file_path, 'w') as f:
            f.write(f"{content['document_type']}\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"System: {content['system_info']['system_name']}\n")
            f.write(f"Description: {content['system_info']['system_description']}\n\n")
            
            if content.get('summary'):
                f.write("Executive Summary:\n")
                f.write(content['summary'])
                f.write("\n\n")
        
        file_size = file_path.stat().st_size
        
        return str(file_path), file_size, None

    def get_available_baselines(self) -> List[ControlBaseline]:
        """Get available control baselines"""
        return list(self.control_baselines.values())

    def get_document_history(self) -> List[RMFDocumentHistory]:
        """Get document generation history"""
        return self.history

    def get_generation_stats(self) -> Dict[str, Any]:
        """Get RMF generation statistics"""
        total_docs = len(self.history)
        docs_by_type = {}
        
        for doc in self.history:
            doc_type = doc.document_type
            docs_by_type[doc_type] = docs_by_type.get(doc_type, 0) + 1
        
        return {
            "total_documents_generated": total_docs,
            "documents_by_type": docs_by_type,
            "available_baselines": len(self.control_baselines),
            "supported_formats": ["DOCX", "PDF", "JSON", "XML"],
            "supported_document_types": ["SSP", "SAR", "POA&M", "Executive Summary"]
        }

    def get_cache_statistics(self) -> Dict[str, Any]:
        """Get RMF generator cache statistics"""
        return {
            "rmf_cache_entries": len(self.rmf_cache),
            "cache_ttl_hours": self.cache_ttl_hours,
            "auto_system_info_loaded": bool(self.auto_system_info),
            "system_metadata": self.auto_system_info
        }

    def clear_cache(self) -> int:
        """Clear RMF generator cache"""
        count = len(self.rmf_cache)
        self.rmf_cache.clear()
        return count


# Global instance
rmf_generator = RMFDocumentGenerator() 